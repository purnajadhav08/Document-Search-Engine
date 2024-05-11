from http.client import HTTPResponse
from django.contrib import messages
from django.shortcuts import render, redirect, get_object_or_404
from .forms import DocumentForm
from .models import Document
import PyPDF2
from docx import Document as DocxDocument
from bson.regex import Regex
import re
import uuid

def preprocess_text(text_content, document_id):
    # Perform preprocessing of the retrieved content
    
    words = re.findall(r'\b\w+\b', text_content.lower())  # Tokenize text into words
    positional_index = {}
    
    for i, word in enumerate(words):
        if word not in positional_index:
            positional_index[word] = {}
        if document_id not in positional_index[word]:
            positional_index[word][document_id] = []
        positional_index[word][document_id].append(i)
    
    return positional_index

# Retrieve main page of the application
def index(request):
    documents = Document.objects.all()
    print("Document Titles:", [document.title for document in documents])  
    return render(request, 'index.html', {'documents': documents})

# View to upload documents, extract content from them and preprocess
def upload_document(request):
    if request.method == 'POST':
        form = DocumentForm(request.POST, request.FILES)
        if form.is_valid():
            for uploaded_file in request.FILES.getlist('file'):
                document = Document(file=uploaded_file)
                document.title = uploaded_file.name  # Set the title to the file name
                document.save()

                # Extract text from the uploaded file
                text_content = extract_text_from_file(uploaded_file)

                if text_content is not None:
                    # Preprocess the text and create positional index
                    positional_index = preprocess_text(text_content, document.id)  # Pass document ID

                    # Save the original content and positional index to the document model
                    document.original_content = text_content
                    # Convert the positional index to the required format
                    positional_index_str = {key: {str(doc_id): positions for doc_id, positions in value.items()} for key, value in positional_index.items()}
                    document.positional_index = positional_index_str

                    document.save()

            #Success message
            messages.success(request, 'Documents uploaded successfully.')
            return redirect('upload_document')
    else:
        form = DocumentForm()
    return render(request, 'upload_document.html', {'form': form})


def extract_text_from_file(uploaded_file):
    text_content = ''
    try:
        if uploaded_file.name.endswith('.pdf'):
            # Handle PDF file
            print("Extracting text from PDF:", uploaded_file.name)
            with uploaded_file.open('rb') as file_object:
                pdf_reader = PyPDF2.PdfReader(file_object)
                for page_num in range(len(pdf_reader.pages)):
                    text_content += pdf_reader.pages[page_num].extract_text()
        elif uploaded_file.name.endswith('.docx'):
            # Handle .docx file using python-docx
            print("Extracting text from docx file:", uploaded_file.name)
            with uploaded_file.open('rb') as file_object:
                docx_doc = DocxDocument(file_object)
                for paragraph in docx_doc.paragraphs:
                    text_content += paragraph.text + '\n'
        elif uploaded_file.name.endswith('.txt'):
            # Handle text file
            print("Extracting text from text file:", uploaded_file.name)
            with uploaded_file.open('r') as file_object:
                text_content = file_object.read()
                # Convert bytes to string if necessary
                if isinstance(text_content, bytes):
                    text_content = text_content.decode('utf-8')
    except Exception as e:
        print("Error extracting text from file:", str(e))
        return None  # Return None if there's an error
    
    print("Extracted text:", text_content)  
    return text_content if text_content is not None else ''  

# View to search documents
def search_document(request):
    if request.method == 'POST':
        query = request.POST.get('query')
        if query:
            return redirect('search_results', query=query)
        else:
            messages.error(request, 'No query provided.')
    return render(request, 'search_document.html')

# View to searcg results
def search_results(request):
    queries = request.GET.get('query', '').split(',')
    queries = [query.strip().lower() for query in queries if query.strip()]
    colors = ['#ffcc00', '#00ffcc', '#cc00ff', '#ff9999', '#99ff99', '#9999ff', '#ff66cc', '#ccff66', '#66ccff', '#ffcc99']

    if not queries:
        return HTTPResponse("No queries provided")

    results = {}
    for query, color in zip(queries, colors):
        matching_documents = Document.objects.filter(original_content__icontains=query)

        for doc in matching_documents:
            original_content = doc.original_content
            highlighted_content = highlight_phrase(original_content, query, color)  

            # If the document is not already in the results dictionary, initialize it
            if doc.id not in results:
                results[doc.id] = {
                    'id': doc.id,
                    'title': doc.title,
                    'highlighted_content': original_content  # Initialize with original content
                }

            # Highlight the current query in the existing highlighted content
            results[doc.id]['highlighted_content'] = highlight_phrase(results[doc.id]['highlighted_content'], query, color)

    final_results = list(results.values())

    return render(request, 'search_results.html', {'results': final_results})


def highlight_phrase(content, query, color):
    pattern = r'\b{}\b'.format(re.escape(query))
    highlighted_content = re.sub(pattern, r'<span style="background-color: {}; color: black;">\g<0></span>'.format(color), content, flags=re.IGNORECASE)
    if content.lower().startswith(query.lower()):
        highlighted_content = re.sub(r'^(?:<span style="background-color: {}; color: yellow;">)?({})'.format(re.escape(color), re.escape(query)), r'<span style="background-color: {}; color: pink;">\1'.format(color), highlighted_content, flags=re.IGNORECASE)
    return highlighted_content

# Display documents details: title and content
def document_detail(request, document_id):
    document = get_object_or_404(Document, pk=document_id)
    return render(request, 'document_detail.html', {'document': document})
